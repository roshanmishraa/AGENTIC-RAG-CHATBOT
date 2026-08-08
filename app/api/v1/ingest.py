from __future__ import annotations
import io
import re
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd

from app.db.session import get_db
from app.db.models import Document, User
from app.db.vectorstore import upsert_chunks
from app.security.rbac import get_current_user
from app.security.rate_limiter import rate_limit_dependency
from app.storage import get_storage          # ← R2 ya Local storage
from app.settings import settings
from app.observability.logger import get_logger

router = APIRouter(prefix="/ingest", tags=["ingest"])
logger = get_logger(__name__)

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024       # 10 MB hard limit

splitter = RecursiveCharacterTextSplitter(
    chunk_size=settings.CHUNK_SIZE,
    chunk_overlap=settings.CHUNK_OVERLAP,
)


# ──────────────────────────────────────────────────────────
# Text cleaning  (PDF artifacts, unicode normalisation, etc.)
# ──────────────────────────────────────────────────────────
def _clean_text(text: str) -> str:
    text = re.sub(r"-\n(\w)", r"\1", text)                           # fix hyphenated line-breaks
    text = re.sub(r"\n{3,}", "\n\n", text)                           # collapse excess blank lines
    text = re.sub(r"(?i)(page\s+\d+\s+of\s+\d+|- \d+ -)", "", text) # strip page-number artifacts
    text = (text.replace("\u2013", "-").replace("\u2014", "--")
                .replace("\u2018", "'").replace("\u2019", "'")
                .replace("\u201c", '"').replace("\u201d", '"'))
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", "", text)
    return text.strip()


# ──────────────────────────────────────────────────────────
# Extractors
# ──────────────────────────────────────────────────────────
def _extract_pdf(file_bytes: bytes) -> list[dict]:
    """Returns [{"text": str, "page_number": int}, ...] — one entry per page."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        raw = page.extract_text() or ""
        text = _clean_text(raw)
        if text:
            pages.append({"text": text, "page_number": i + 1})
    return pages


def _extract_docx(file_bytes: bytes) -> list[dict]:
    doc = DocxDocument(io.BytesIO(file_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [{"text": _clean_text(full_text), "page_number": None}]


def _extract_csv(file_bytes: bytes) -> list[dict]:
    df = pd.read_csv(io.BytesIO(file_bytes))
    text = df.to_string(index=False)
    return [{"text": text, "page_number": None}]


EXTRACTORS = {
    "pdf":  _extract_pdf,
    "docx": _extract_docx,
    "csv":  _extract_csv,
}


# ──────────────────────────────────────────────────────────
# Upload endpoint
# ──────────────────────────────────────────────────────────
@router.post("/upload", dependencies=[Depends(rate_limit_dependency)])
async def upload_document(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # ── 1. Validate file type ──────────────────────────────
    file_ext = (file.filename or "").split(".")[-1].lower()
    if file_ext not in EXTRACTORS:
        raise HTTPException(
            400,
            f"Unsupported file type: .{file_ext}. "
            f"Supported: {list(EXTRACTORS.keys())}",
        )

    # ── 2. Read bytes + size check ─────────────────────────
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(400, "Uploaded file is empty.")
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            413,
            f"File too large ({len(file_bytes) // (1024*1024)} MB). "
            f"Maximum allowed size is 10 MB.",
        )

    # ── 3. Save original file to storage (R2 or local) ─────
    storage = get_storage()
    try:
        storage_path = await storage.save(
            user_id=user.id,
            filename=file.filename or f"upload.{file_ext}",
            data=file_bytes,
        )
    except Exception as exc:
        logger.error(f"Storage save failed: {exc}", extra={"user_id": user.id})
        raise HTTPException(500, "Failed to save file to storage.")

    # ── 4. Create Document row (status=processing) ─────────
    document = Document(
        owner_id=user.id,
        filename=file.filename,
        file_type=file_ext,
        status="processing",
        storage_path=storage_path,   # ← local path or R2 key
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    # ── 5. Parse → chunk → embed → store ──────────────────
    try:
        pages = EXTRACTORS[file_ext](file_bytes)
        if not pages:
            raise ValueError("No extractable text found in file.")

        chunks: list[dict] = []
        chunk_index = 0
        for page in pages:
            for chunk_text in splitter.split_text(page["text"]):
                chunks.append({
                    "content": chunk_text,
                    "chunk_index": chunk_index,
                    "page_number": page["page_number"],
                })
                chunk_index += 1

        # owner_id pass karna zaroori hai — Pinecone metadata mein
        # user isolation ke liye (audit report CRIT-3 fix)
        await upsert_chunks(
            db,
            document_id=document.id,
            chunks=chunks,
            owner_id=user.id,        # ← CRIT-3 fix
        )

        document.status = "ready"
        await db.commit()

        logger.info(
            f"Document ingested: {file.filename} ({len(chunks)} chunks)",
            extra={"user_id": user.id, "document_id": document.id},
        )

        return {
            "document_id": document.id,
            "filename": document.filename,
            "status": "ready",
            "chunks_created": len(chunks),
            "storage_path": storage_path,
        }

    except HTTPException:
        raise

    except Exception as exc:
        # Mark failed but keep the Document row so user knows upload happened
        document.status = "failed"
        await db.commit()

        # Try to clean up the stored file — don't crash if cleanup also fails
        try:
            await storage.delete(storage_path)
        except Exception:
            pass

        logger.error(
            f"Ingestion failed for {file.filename}: {exc}",
            extra={"user_id": user.id, "document_id": document.id},
        )
        raise HTTPException(500, f"Failed to process document: {str(exc)}")


# ──────────────────────────────────────────────────────────
# List documents  (user sirf apne docs dekh sakta hai)
# ──────────────────────────────────────────────────────────
@router.get("/documents")
async def list_documents(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Document).where(Document.owner_id == user.id)
        .order_by(Document.uploaded_at.desc())
    )
    docs = result.scalars().all()
    return [
        {
            "id": d.id,
            "filename": d.filename,
            "file_type": d.file_type,
            "status": d.status,
            "uploaded_at": d.uploaded_at,
            "storage_path": d.storage_path,
        }
        for d in docs
    ]


# ──────────────────────────────────────────────────────────
# Delete document  (owner hi delete kar sakta hai)
# ──────────────────────────────────────────────────────────
@router.delete("/documents/{document_id}", status_code=204)
async def delete_document(
    document_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.owner_id == user.id,   # ← ownership check
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found or does not belong to you.")

    # Storage se file delete karo
    if doc.storage_path:
        try:
            storage = get_storage()
            await storage.delete(doc.storage_path)
        except Exception as exc:
            logger.warning(
                f"Storage delete failed for {doc.storage_path}: {exc}",
                extra={"user_id": user.id, "document_id": document_id},
            )

    # DB se document + chunks delete (cascade handle karta hai chunks)
    await db.delete(doc)
    await db.commit()